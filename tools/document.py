"""
tools/document.py — Document parsing helpers for PDF and Word files.
"""

from __future__ import annotations

import json
import math
import os
import re
from dataclasses import dataclass
from typing import Iterable

DEFAULT_MAX_CHARS = 14000
DEFAULT_CHUNK_SIZE = 12000
MAX_CHUNK_SIZE = 50000
MAX_QUERY_MATCHES = 8

SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx"}


@dataclass
class TextSegment:
    """A readable unit of document text with optional location metadata."""

    text: str
    label: str
    page: int | None = None
    paragraph: int | None = None


def _json(data: dict) -> str:
    return json.dumps(data, ensure_ascii=False)


def _positive_int(value: int | str | None, default: int, minimum: int = 1, maximum: int | None = None) -> int:
    try:
        parsed = int(value) if value is not None else default
    except (TypeError, ValueError):
        parsed = default
    parsed = max(minimum, parsed)
    if maximum is not None:
        parsed = min(parsed, maximum)
    return parsed


def _parse_page_spec(pages: str | None, page_count: int) -> list[int]:
    """Parse a 1-based page specification like "1-3,8" into 0-based indexes."""
    if not pages:
        return []

    if len(pages) > 1000:
        raise ValueError("Page specification exceeds the 1000-character limit")
    parts = pages.split(",")
    if len(parts) > 200:
        raise ValueError("Page specification contains too many ranges")
    selected: set[int] = set()
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start_raw, end_raw = part.split("-", 1)
            try:
                start = int(start_raw)
                end = int(end_raw)
            except ValueError as exc:
                raise ValueError(f"Invalid page range: {part}") from exc
            if start > end:
                start, end = end, start
            bounded_start = max(1, start)
            bounded_end = min(page_count, end)
            if bounded_start <= bounded_end:
                selected.update(range(bounded_start - 1, bounded_end))
        else:
            try:
                page = int(part)
            except ValueError as exc:
                raise ValueError(f"Invalid page number: {part}") from exc
            if 1 <= page <= page_count:
                selected.add(page - 1)

    if not selected:
        raise ValueError(f"No pages from '{pages}' are inside this document's 1-{page_count} page range")
    return sorted(selected)


def _chunk_text(text: str, chunk_size: int = DEFAULT_CHUNK_SIZE) -> list[str]:
    chunk_size = _positive_int(chunk_size, DEFAULT_CHUNK_SIZE, minimum=1000, maximum=MAX_CHUNK_SIZE)
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        hard_end = min(len(text), start + chunk_size)
        end = hard_end
        if hard_end < len(text):
            window = text[start:hard_end]
            boundary = max(window.rfind("\n\n"), window.rfind("\n"), window.rfind(". "))
            if boundary >= chunk_size // 2:
                end = start + boundary + 1
        chunks.append(text[start:end])
        start = end
    return chunks


def _snippet(text: str, query: str, max_chars: int = 900) -> str:
    lower = text.lower()
    query_lower = query.lower().strip()
    pos = lower.find(query_lower) if query_lower else -1
    if pos < 0:
        terms = [term for term in re.findall(r"\w+", query_lower) if len(term) > 2]
        positions = [lower.find(term) for term in terms if lower.find(term) >= 0]
        pos = min(positions) if positions else 0

    start = max(0, pos - max_chars // 3)
    end = min(len(text), start + max_chars)
    start = max(0, end - max_chars)
    result = text[start:end].strip()
    if start:
        result = "..." + result
    if end < len(text):
        result += "..."
    return result


def _score_segment(text: str, query: str) -> int:
    lower = text.lower()
    query_lower = query.lower().strip()
    terms = [term for term in re.findall(r"\w+", query_lower) if len(term) > 2]
    score = lower.count(query_lower) * 8 if query_lower else 0
    score += sum(lower.count(term) for term in terms)
    return score


def _search_segments(segments: Iterable[TextSegment], query: str, max_matches: int = MAX_QUERY_MATCHES) -> list[dict]:
    matches = []
    for segment in segments:
        score = _score_segment(segment.text, query)
        if score <= 0:
            continue
        match = {
            "label": segment.label,
            "score": score,
            "snippet": _snippet(segment.text, query),
        }
        if segment.page is not None:
            match["page"] = segment.page
        if segment.paragraph is not None:
            match["paragraph"] = segment.paragraph
        matches.append(match)

    matches.sort(key=lambda item: item["score"], reverse=True)
    return matches[:max_matches]


def _join_segments(segments: Iterable[TextSegment]) -> str:
    parts = []
    for segment in segments:
        text = segment.text.strip()
        if not text:
            continue
        parts.append(f"[{segment.label}]\n{text}")
    return "\n\n".join(parts)


def _extract_pdf_segments(file_path: str, pages: str | None = None, preview_chars: int | None = None) -> tuple[list[TextSegment], dict]:
    try:
        import pypdf
    except ImportError as exc:
        raise RuntimeError("Missing required dependency. Please run: pip install pypdf") from exc

    segments: list[TextSegment] = []
    preview_limited = False
    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        if reader.is_encrypted:
            try:
                if not reader.decrypt(""):
                    raise ValueError("password required")
            except Exception as exc:
                raise ValueError("PDF is encrypted and cannot be read without a password") from exc
        page_count = len(reader.pages)
        selected_pages = _parse_page_spec(pages, page_count) if pages else range(page_count)
        extracted_chars = 0

        for index in selected_pages:
            page_num = index + 1
            try:
                text = reader.pages[index].extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                segments.append(TextSegment(text=text.strip(), label=f"page {page_num}", page=page_num))
                extracted_chars += len(text)
            if preview_chars and extracted_chars >= preview_chars:
                preview_limited = page_num < page_count
                break

    info = {
        "document_type": "pdf",
        "page_count": page_count,
        "pages_requested": pages,
        "pages_extracted": [segment.page for segment in segments if segment.page is not None],
        "preview_limited": preview_limited,
    }
    return segments, info


def _iter_docx_blocks(doc) -> Iterable[str]:
    if hasattr(doc, "iter_inner_content"):
        for block in doc.iter_inner_content():
            if hasattr(block, "rows"):
                for row in block.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        yield " | ".join(cells)
            else:
                text = block.text.strip()
                if text:
                    yield text
        return
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                yield " | ".join(cells)


def _extract_docx_segments(file_path: str, preview_chars: int | None = None) -> tuple[list[TextSegment], dict]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("Missing required dependency. Please run: pip install python-docx") from exc

    document = docx.Document(file_path)
    segments: list[TextSegment] = []
    extracted_chars = 0
    preview_limited = False
    for block_index, text in enumerate(_iter_docx_blocks(document), start=1):
        segments.append(TextSegment(text=text, label=f"block {block_index}", paragraph=block_index))
        extracted_chars += len(text)
        if preview_chars and extracted_chars >= preview_chars:
            preview_limited = True
            break

    info = {
        "document_type": "docx",
        "block_count_extracted": len(segments),
        "pages_requested": None,
        "preview_limited": preview_limited,
    }
    return segments, info


def extract_document_text(file_path: str, pages: str | None = None) -> tuple[str, dict]:
    """
    Return full extracted document text plus metadata for indexing callers.
    
    This function acts as a unified entry point for reading both PDF and DOCX files.
    It extracts the text segments and joins them into a single continuous string
    suitable for search indexing, while also returning metadata like page counts.
    
    Args:
        file_path (str): The absolute or relative path to the document file.
        pages (str | None): An optional page range string (e.g., "1-5") for PDFs.
            
    Returns:
        tuple[str, dict]: A tuple where the first element is the extracted text 
            string and the second is a dictionary of document metadata.
            
    Raises:
        ValueError: If the document type is unsupported.
    """
    # Determine the file extension to route to the correct extractor
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        # Extract segments from PDF with optional page range
        segments, info = _extract_pdf_segments(file_path, pages=pages)
    elif ext == ".docx":
        # Extract blocks from DOCX (page ranges are not supported)
        segments, info = _extract_docx_segments(file_path)
    else:
        raise ValueError(f"Unsupported document type: {ext}")
        
    # Join the extracted segments into a single readable string
    text = _join_segments(segments)
    info["char_count"] = len(text)
    
    return text, info


def read_document(
    file_path: str,
    pages: str | None = None,
    query: str | None = None,
    chunk: int | str | None = None,
    chunk_size: int | str = DEFAULT_CHUNK_SIZE,
    max_chars: int | str = DEFAULT_MAX_CHARS,
) -> str:
    """
    Extract text from a PDF or Word document with large-file navigation capabilities.

    This tool enables reading content from binary document formats (.pdf, .docx).
    To handle large files, it provides features like page filtering, semantic chunking,
    and query-based snippet search.

    Args:
        file_path (str): Absolute or relative path to a .pdf or .docx file.
        pages (str | None): Optional 1-based PDF pages/ranges, such as "1-3,8".
            Only applies to PDFs.
        query (str | None): Optional search query; if provided, returns matched
            text snippets instead of a whole document dump.
        chunk (int | str | None): Optional 0-based chunk number to retrieve
            when navigating large extracted text sequentially.
        chunk_size (int | str): Approximate characters per chunk (default is 12000).
        max_chars (int | str): Maximum characters returned in preview/text fields.
            Used to prevent overflowing the LLM context window.
            
    Returns:
        str: A JSON-encoded string containing the extracted text, search matches,
             or metadata about the document. Contains an 'error' key if reading fails.
    """
    if not os.path.exists(file_path):
        return _json({"error": f"File not found: {file_path}"})
    if not os.path.isfile(file_path):
        return _json({"error": f"Not a file: {file_path}"})

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_DOCUMENT_EXTENSIONS:
        return _json({
            "error": f"Unsupported file type: {ext}. Use read_file for text files; read_document supports .pdf and .docx.",
            "supported_extensions": sorted(SUPPORTED_DOCUMENT_EXTENSIONS),
        })

    max_chars_int = _positive_int(max_chars, DEFAULT_MAX_CHARS, minimum=2000, maximum=MAX_CHUNK_SIZE)
    chunk_size_int = _positive_int(chunk_size, DEFAULT_CHUNK_SIZE, minimum=1000, maximum=MAX_CHUNK_SIZE)
    if query is not None:
        query = str(query).strip()
        if len(query) > 4000:
            return _json({"error": "query exceeds the 4000-character limit"})
    chunk_index = None
    if chunk is not None:
        chunk_index = _positive_int(chunk, 0, minimum=0)

    try:
        preview_only = not query and chunk_index is None and not pages
        preview_chars = max_chars_int if preview_only else None
        if ext == ".pdf":
            segments, info = _extract_pdf_segments(file_path, pages=pages, preview_chars=preview_chars)
        else:
            segments, info = _extract_docx_segments(file_path, preview_chars=preview_chars)
            if pages:
                info["warning"] = "The pages parameter is only supported for PDFs; extracted DOCX text by document blocks instead."

        base = {
            "file": file_path,
            "size_bytes": os.path.getsize(file_path),
            **info,
        }

        if query:
            matches = _search_segments(segments, query)
            base.update({
                "mode": "query",
                "query": query,
                "matches": matches,
                "match_count": len(matches),
                "guidance": "Use pages for exact PDF pages, or chunk for surrounding document context if a match needs more detail.",
            })
            return _json(base)

        text = _join_segments(segments).strip()
        chunks = _chunk_text(text, chunk_size=chunk_size_int) if text else [""]
        total_chunks = len(chunks)
        selected_chunk = 0 if chunk_index is None else min(chunk_index, total_chunks - 1)
        selected_text = chunks[selected_chunk]
        truncated = len(selected_text) > max_chars_int
        if truncated:
            selected_text = selected_text[:max_chars_int].rstrip() + "\n\n...[Chunk truncated. Request a smaller chunk_size or a more specific page/query.]"

        base.update({
            "mode": "preview" if preview_only and total_chunks <= 1 else "chunk",
            "text": selected_text,
            "returned_chars": len(selected_text),
            "total_extracted_chars": len(text),
            "truncated": truncated or total_chunks > 1 or bool(info.get("preview_limited")),
            "navigation": {
                "chunk": selected_chunk,
                "chunk_size": chunk_size_int,
                "total_chunks": total_chunks,
                "next_chunk": selected_chunk + 1 if selected_chunk + 1 < total_chunks else None,
                "previous_chunk": selected_chunk - 1 if selected_chunk > 0 else None,
                "pages_parameter": "For PDFs, pass pages like '4-6' or '1,3,9'.",
                "query_parameter": "Pass query to get the most relevant snippets without reading the whole document.",
            },
        })

        if preview_only and ext == ".pdf" and info.get("page_count") and info.get("pages_extracted"):
            extracted = info["pages_extracted"]
            last_page = extracted[-1] if extracted else None
            if last_page and last_page < info["page_count"]:
                base["guidance"] = f"Preview stopped at page {last_page}. Request later pages with pages='{last_page + 1}-{min(last_page + 5, info['page_count'])}'."

        if chunk_index is not None and chunk_index >= total_chunks:
            base["warning"] = f"Requested chunk {chunk_index}, but only {total_chunks} chunk(s) exist; returned the last chunk."

        if text and total_chunks > 1:
            base["estimated_total_chunks"] = math.ceil(len(text) / chunk_size_int)
        return _json(base)

    except ValueError as exc:
        return _json({"error": str(exc)})
    except RuntimeError as exc:
        return _json({"error": str(exc)})
    except Exception as exc:
        return _json({"error": f"Error reading document: {exc}"})
